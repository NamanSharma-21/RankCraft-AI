"""
Resume Parser Module
Extracts text from PDF, DOCX, and TXT files reliably and extracts candidate metadata.
"""

import io
import os
import re
from typing import Tuple, Optional, Union
import pypdf
from docx import Document


class FileExtractionError(Exception):
    """Raised when text cannot be extracted from a resume file."""
    pass


class UnsupportedFormatError(Exception):
    """Raised when an unsupported file format is uploaded."""
    pass


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_text_from_pdf(file_input: Union[str, bytes, io.BytesIO]) -> str:
    """
    Extracts text from a PDF file using pypdf.
    Accepts file path, raw bytes, or BytesIO stream.
    """
    try:
        if isinstance(file_input, (bytes, bytearray)):
            stream = io.BytesIO(file_input)
        elif isinstance(file_input, str):
            if not os.path.exists(file_input):
                raise FileExtractionError(f"PDF file not found: {file_input}")
            stream = open(file_input, "rb")
        else:
            stream = file_input

        reader = pypdf.PdfReader(stream)
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as e:
                raise FileExtractionError(f"PDF is encrypted and cannot be read: {e}")

        text_pages = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_pages.append(page_text)

        full_text = "\n".join(text_pages).strip()
        if not full_text:
            raise FileExtractionError("Extracted PDF content is empty or contains non-extractable text.")
        return full_text
    except Exception as e:
        if isinstance(e, FileExtractionError):
            raise
        raise FileExtractionError(f"Failed to extract PDF text: {str(e)}")


def extract_text_from_docx(file_input: Union[str, bytes, io.BytesIO]) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    Accepts file path, raw bytes, or BytesIO stream.
    """
    try:
        if isinstance(file_input, (bytes, bytearray)):
            stream = io.BytesIO(file_input)
        elif isinstance(file_input, str):
            if not os.path.exists(file_input):
                raise FileExtractionError(f"DOCX file not found: {file_input}")
            stream = open(file_input, "rb")
        else:
            stream = file_input

        doc = Document(stream)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract table text if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        full_text = "\n".join(paragraphs).strip()
        if not full_text:
            raise FileExtractionError("Extracted DOCX content is empty.")
        return full_text
    except Exception as e:
        if isinstance(e, FileExtractionError):
            raise
        raise FileExtractionError(f"Failed to extract DOCX text: {str(e)}")


def extract_text_from_txt(file_input: Union[str, bytes, io.BytesIO]) -> str:
    """
    Extracts text from a plain TXT file with utf-8 / fallback encodings.
    """
    try:
        if isinstance(file_input, (bytes, bytearray)):
            raw_bytes = file_input
        elif isinstance(file_input, str):
            if not os.path.exists(file_input):
                raise FileExtractionError(f"TXT file not found: {file_input}")
            with open(file_input, "rb") as f:
                raw_bytes = f.read()
        else:
            raw_bytes = file_input.read()

        for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                text = raw_bytes.decode(encoding).strip()
                if text:
                    return text
            except (UnicodeDecodeError, AttributeError):
                continue

        raise FileExtractionError("Could not decode TXT file with standard encodings.")
    except Exception as e:
        if isinstance(e, FileExtractionError):
            raise
        raise FileExtractionError(f"Failed to read TXT text: {str(e)}")


def extract_candidate_name(raw_text: str, filename: Optional[str] = None) -> str:
    """
    Heuristically extracts candidate name from the top lines of resume text,
    falling back to formatted filename if no name pattern is found.
    """
    if raw_text:
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        for line in lines[:5]:
            # Skip email lines, phone numbers, urls, or section headers
            if "@" in line or "http" in line or "resume" in line.lower() or "curriculum" in line.lower():
                continue
            # Look for 2 to 4 words starting with capital letters, length < 40
            words = line.split()
            if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w[0].isalpha()):
                # Strip titles like CPA, PhD, Dr.
                cleaned = re.sub(r"[,|].*", "", line).strip()
                return cleaned

    # Fallback to filename
    if filename:
        clean_name = os.path.splitext(os.path.basename(filename))[0]
        clean_name = re.sub(r"^candidate_\d+_", "", clean_name, flags=re.IGNORECASE)
        clean_name = re.sub(r"[_\-\.]+", " ", clean_name).title()
        return clean_name

    return "Candidate"


def parse_resume(file_input: Union[str, bytes, io.BytesIO], filename: str) -> dict:
    """
    Unified entry point for resume parsing.
    Returns dictionary with extracted text, candidate name, format, and metadata.
    """
    _, ext = os.path.splitext(filename.lower())
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFormatError(
            f"Unsupported file format '{ext}'. Supported formats are: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    if ext == ".pdf":
        raw_text = extract_text_from_pdf(file_input)
        file_type = "PDF"
    elif ext == ".docx":
        raw_text = extract_text_from_docx(file_input)
        file_type = "DOCX"
    elif ext == ".txt":
        raw_text = extract_text_from_txt(file_input)
        file_type = "TXT"
    else:
        raise UnsupportedFormatError(f"Unhandled file extension: {ext}")

    candidate_name = extract_candidate_name(raw_text, filename)
    snippet = raw_text[:250].replace("\n", " ").strip() + ("..." if len(raw_text) > 250 else "")

    return {
        "file_name": filename,
        "file_type": file_type,
        "candidate_name": candidate_name,
        "raw_text": raw_text,
        "character_count": len(raw_text),
        "word_count": len(raw_text.split()),
        "snippet": snippet
    }
