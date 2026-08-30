"""
Sample Data Generator for AI Resume Screening System
Generates realistic demonstration resumes in PDF, DOCX, and TXT formats,
along with job descriptions, candidate metadata, and evaluation ground-truth labels.
"""

import os
import csv
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# Directories
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(BASE_DIR, "data")
RESUMES_DIR = os.path.join(DATA_DIR, "resumes")
JD_DIR = os.path.join(DATA_DIR, "job_descriptions")
METADATA_DIR = os.path.join(DATA_DIR, "metadata")
EVAL_DIR = os.path.join(DATA_DIR, "evaluation")

for d in [RESUMES_DIR, JD_DIR, METADATA_DIR, EVAL_DIR]:
    os.makedirs(d, exist_ok=True)

# -------------------------------------------------------------
# 1. Job Descriptions
# -------------------------------------------------------------
JOB_DESCRIPTIONS = {
    "job_01_senior_ai_ml_engineer.txt": {
        "job_id": "JOB-01",
        "title": "Senior AI / Machine Learning Engineer",
        "content": """Job Title: Senior AI / Machine Learning Engineer
Department: AI & Data Science
Location: Remote / Hybrid

About the Role:
We are seeking an experienced Senior AI/Machine Learning Engineer to lead the design and deployment of production ML systems. You will work on cutting-edge Natural Language Processing (NLP), Large Language Model (LLM) fine-tuning, and scalable inference pipelines.

Key Responsibilities:
- Design, train, and deploy machine learning and deep learning models for production NLP applications.
- Build feature extraction pipelines and implement statistical and neural ranking systems (TF-IDF, embeddings, transformer architectures).
- Develop robust REST APIs using Python and FastAPI for low-latency model serving.
- Implement end-to-end MLOps workflows using Docker, Kubernetes, and CI/CD pipelines.
- Collaborate with cloud engineers to optimize model inference on AWS infrastructure.
- Write clean, modular, well-tested Python code adhering to strict software engineering standards.

Requirements & Qualifications:
- 5+ years of software development experience with strong proficiency in Python.
- Deep expertise in ML frameworks: PyTorch, TensorFlow, Scikit-Learn, Hugging Face Transformers.
- Strong background in Natural Language Processing (NLP), text preprocessing, tokenization, and vectorization.
- Experience with FastAPI, Flask, or Django for backend API development.
- Hands-on experience with Docker, Kubernetes, Git, CI/CD, and AWS cloud services (S3, EC2, SageMaker).
- Solid understanding of SQL, relational databases (PostgreSQL), and vector databases.
- Familiarity with agile development, unit testing (pytest), and reproducible data pipelines.
"""
    },
    "job_02_fullstack_python_developer.txt": {
        "job_id": "JOB-02",
        "title": "Full-Stack Python Developer",
        "content": """Job Title: Full-Stack Python Developer
Department: Engineering
Location: Hybrid

About the Role:
We are looking for a skilled Full-Stack Python Developer to build modern, high-performance web applications. You will be responsible for creating robust backend services and responsive user interfaces.

Key Responsibilities:
- Develop scalable backend RESTful APIs using Python, FastAPI, and Django.
- Design database schemas and write optimized queries with PostgreSQL and SQLAlchemy.
- Build responsive, accessible frontend interfaces using HTML5, CSS3, JavaScript, and modern UI patterns.
- Containerize applications with Docker and deploy services to AWS or cloud environments.
- Maintain test coverage through comprehensive unit and integration testing (pytest, Jest).
- Collaborate with product designers and engineering teams in an Agile sprint environment.

Requirements & Qualifications:
- 3+ years of professional full-stack web development experience.
- Strong proficiency in Python, FastAPI, Django, or Flask.
- Solid front-end skills in JavaScript, HTML5, CSS3, DOM manipulation, and responsive design.
- Experience with relational databases: PostgreSQL, MySQL, and database migrations.
- Experience with Docker, Git version control, and GitHub Actions CI/CD workflows.
- Strong understanding of RESTful API architecture, HTTP protocols, and JSON data exchange.
"""
    },
    "job_03_data_analyst.txt": {
        "job_id": "JOB-03",
        "title": "Data Analyst & Business Intelligence Specialist",
        "content": """Job Title: Data Analyst & Business Intelligence Specialist
Department: Analytics & BI
Location: On-site / Hybrid

About the Role:
We are looking for a Data Analyst to transform complex datasets into actionable business intelligence dashboards and analytical reports.

Key Responsibilities:
- Extract, clean, and transform data from diverse SQL databases and data warehouses.
- Build interactive dashboards and reporting suites in Tableau and Power BI.
- Perform exploratory data analysis and statistical modeling using Python, Pandas, and NumPy.
- Automate recurring business reporting and ETL data pipelines.
- Present data-driven findings and insights to executive stakeholders.

Requirements & Qualifications:
- 3+ years of experience in data analysis, business intelligence, or quantitative analytics.
- Advanced SQL proficiency (complex joins, CTEs, window functions, query optimization).
- Strong data manipulation skills in Python (Pandas, NumPy) and Excel (VLOOKUP, Pivot Tables).
- Proven track record building enterprise dashboards in Tableau or Power BI.
- Understanding of ETL processes, data warehousing concepts, and data quality assurance.
- Strong communication and data storytelling skills.
"""
    }
}

# -------------------------------------------------------------
# 2. Candidate Resumes Data
# -------------------------------------------------------------
CANDIDATES = [
    {
        "candidate_id": "CAND-001",
        "name": "Dr. Alex Rivera",
        "file_name": "candidate_01_senior_ml_engineer.pdf",
        "file_type": "PDF",
        "primary_domain": "AI / Machine Learning",
        "years_experience": 6,
        "text": """Dr. Alex Rivera
Senior Machine Learning Engineer | NLP Specialist
Email: alex.rivera.ai@example.com | GitHub: github.com/alexrivera-ml | Location: San Francisco, CA

PROFESSIONAL SUMMARY:
Senior AI/ML Engineer with 6+ years of experience architecting and deploying production NLP and deep learning systems. Proven expertise in Natural Language Processing, Transformer models, PyTorch, TensorFlow, Scikit-Learn, and low-latency FastAPI model serving. Experienced in MLOps, Docker containerization, Kubernetes orchestration, and AWS cloud pipelines.

TECHNICAL SKILLS:
- Languages & Frameworks: Python, PyTorch, TensorFlow, Scikit-Learn, Hugging Face, NumPy, Pandas
- NLP & Information Retrieval: TF-IDF, Word Embeddings, Transformers, BERT, LLM Fine-tuning, Tokenization, Text Preprocessing
- Backend & APIs: FastAPI, Flask, REST APIs, Pydantic, Uvicorn, Asynchronous Programming
- Cloud & MLOps: AWS (S3, EC2, SageMaker), Docker, Kubernetes, CI/CD, MLflow, Git
- Databases & Tools: PostgreSQL, SQL, Vector Databases, Pytest, Linux, Agile/Scrum

WORK EXPERIENCE:
Senior Machine Learning Engineer | DeepData Labs (2021 – Present)
- Designed and scaled transformer-based NLP pipelines serving 10M+ daily inference queries.
- Built low-latency REST APIs using Python and FastAPI for real-time document classification and ranking.
- Implemented automated CI/CD and Dockerized deployment workflows on AWS Kubernetes clusters.
- Fine-tuned open-source LLMs and built TF-IDF baseline search systems with high precision and recall.

Machine Learning Engineer | Apex Analytics (2018 – 2021)
- Developed text preprocessing, tokenization, and vectorization modules using Scikit-Learn and NLTK.
- Deployed regression and classification models using Python, Docker, and PostgreSQL.
- Conducted unit testing with pytest and maintained 90%+ code coverage.

EDUCATION:
- Ph.D. in Computer Science (Machine Learning Focus), Stanford University (2018)
- B.S. in Computer Science & Applied Mathematics, UC Berkeley (2014)
"""
    },
    {
        "candidate_id": "CAND-002",
        "name": "Sarah Chen",
        "file_name": "candidate_02_nlp_data_scientist.docx",
        "file_type": "DOCX",
        "primary_domain": "NLP & Data Science",
        "years_experience": 4,
        "text": """Sarah Chen
Machine Learning Scientist & NLP Researcher
Email: sarah.chen.nlp@example.com | Portfolio: sarahchen-ai.example.org | Location: Boston, MA

SUMMARY:
Data Scientist and NLP Specialist with 4 years of experience building statistical language models, text vectorization systems, and machine learning pipelines. Strong background in Python, Scikit-Learn, PyTorch, TF-IDF ranking, and cloud deployment.

CORE COMPETENCIES:
- Machine Learning & NLP: Python, PyTorch, Scikit-Learn, Hugging Face, NLTK, TF-IDF, Cosine Similarity, Tokenization
- Data Science & Analytics: Pandas, NumPy, Scipy, Matplotlib, Seaborn, Feature Engineering
- Software Engineering & Deployment: FastAPI, Flask, Docker, Git, CI/CD, REST APIs, Pytest
- Databases & Cloud: PostgreSQL, SQL queries, AWS S3, Linux environment

PROFESSIONAL EXPERIENCE:
NLP Data Scientist | TextFlow AI (2022 – Present)
- Developed text preprocessing and TF-IDF feature extraction pipelines for automated resume and job matching.
- Implemented cosine similarity algorithms and semantic ranking models evaluated on precision and NDCG metrics.
- Built FastAPI microservices containerized with Docker for real-time scoring.

Junior ML Engineer | Cognition Corp (2020 – 2022)
- Extracted and normalized unstructured text from PDF and DOCX documents using Python libraries.
- Implemented supervised and unsupervised models using Scikit-Learn and PyTorch.
- Automated data validation and unit testing using pytest and GitHub Actions.

EDUCATION:
- M.S. in Data Science, Massachusetts Institute of Technology (MIT), 2020
- B.S. in Statistics and Computer Science, University of Michigan, 2018
"""
    },
    {
        "candidate_id": "CAND-003",
        "name": "Marcus Vance",
        "file_name": "candidate_03_fullstack_python_dev.pdf",
        "file_type": "PDF",
        "primary_domain": "Full-Stack Web Development",
        "years_experience": 5,
        "text": """Marcus Vance
Lead Full-Stack Python Developer
Email: marcus.vance.dev@example.com | GitHub: github.com/marcusvance | Location: Austin, TX

PROFESSIONAL SUMMARY:
Versatile Full-Stack Python Developer with 5+ years of experience building high-throughput web applications, RESTful APIs, and responsive frontends. Deep expertise in Python, FastAPI, Django, PostgreSQL, HTML5, CSS3, JavaScript, Docker, and AWS cloud infrastructure.

TECHNICAL SKILLS:
- Backend: Python, FastAPI, Django, Flask, SQLAlchemy, Celery, RESTful API design
- Frontend: JavaScript (ES6+), HTML5, CSS3, Responsive Design, DOM manipulation, UI/UX
- Database: PostgreSQL, MySQL, Redis, Database migrations, SQL query optimization
- DevOps & Cloud: Docker, AWS (EC2, S3, RDS), Git, GitHub Actions CI/CD, Nginx, Linux
- Testing & Methodologies: Pytest, Test-Driven Development (TDD), Agile/Scrum

EXPERIENCE:
Senior Full-Stack Developer | CloudPeak Systems (2021 – Present)
- Architected and deployed scalable REST APIs in Python using FastAPI and Django.
- Built responsive, accessible user interfaces using semantic HTML5, modern CSS3, and JavaScript.
- Managed PostgreSQL database schema design, index tuning, and data migration scripts.
- Configured CI/CD deployment pipelines using Docker containers and AWS EC2/S3.

Full-Stack Software Engineer | ByteWorks Inc (2019 – 2021)
- Developed backend microservices in Python with Flask and PostgreSQL.
- Implemented interactive front-end dashboards using vanilla JavaScript, HTML, and CSS.
- Wrote extensive unit tests in pytest and integration tests to ensure system reliability.

EDUCATION:
- B.S. in Software Engineering, University of Texas at Austin (2019)
"""
    },
    {
        "candidate_id": "CAND-004",
        "name": "Elena Rostova",
        "file_name": "candidate_04_backend_python_engineer.txt",
        "file_type": "TXT",
        "primary_domain": "Backend Engineering",
        "years_experience": 4,
        "text": """Elena Rostova
Backend Python Engineer & API Developer
Email: elena.rostova.eng@example.com | Location: Seattle, WA

PROFESSIONAL SUMMARY:
Backend Software Engineer with 4 years of experience specializing in Python, FastAPI, Django, and database architecture. Proven ability to create resilient microservices, integrate relational databases, and maintain CI/CD pipelines in cloud environments.

SKILLS:
- Programming: Python, SQL, Bash
- Frameworks: FastAPI, Django, Flask, SQLAlchemy, Pydantic, Uvicorn
- Databases: PostgreSQL, SQLite, Redis, Relational Database Modeling
- Infrastructure: Docker, Git, GitHub Actions, CI/CD, AWS, Linux
- Web Standards: RESTful APIs, HTTP, JSON, WebSockets, OpenAPI

PROFESSIONAL EXPERIENCE:
Backend Engineer | Cascade Data Systems (2022 – Present)
- Engineered high-performance backend endpoints using FastAPI and PostgreSQL.
- Implemented robust input validation, serialization, and error handling mechanisms.
- Automated test suites using Pytest with 92% code coverage.
- Containerized applications using Docker and deployed via GitHub Actions CI/CD.

Python Developer | InnovateTech (2020 – 2022)
- Built internal REST APIs and ETL data transformation scripts using Python.
- Designed relational schemas and optimized SQL queries for high-volume transactions.

EDUCATION:
- B.S. in Computer Science, University of Washington (2020)
"""
    },
    {
        "candidate_id": "CAND-005",
        "name": "David Kim",
        "file_name": "candidate_05_senior_data_analyst.pdf",
        "file_type": "PDF",
        "primary_domain": "Data Analytics & BI",
        "years_experience": 5,
        "text": """David Kim
Senior Data Analyst & BI Specialist
Email: david.kim.analytics@example.com | Location: Chicago, IL

PROFESSIONAL SUMMARY:
Analytical and detail-oriented Senior Data Analyst with 5+ years of experience delivering actionable business intelligence, KPI reporting, and statistical data modeling. Expert in advanced SQL, Tableau, Power BI, Python (Pandas, NumPy), and ETL data pipelines.

CORE COMPETENCIES:
- Data Analysis & Modeling: SQL, Python (Pandas, NumPy, Matplotlib), Statistical Analysis, Hypothesis Testing
- Business Intelligence & Visualization: Tableau, Power BI, Excel (Advanced, VBA), Looker
- Data Engineering & ETL: Data Warehousing, Data Cleaning, Data Transformation, PostgreSQL
- Communication: Executive Presentation, Stakeholder Management, Data Storytelling

EXPERIENCE:
Senior Business Intelligence Analyst | Horizon Financial (2021 – Present)
- Developed interactive executive dashboards in Tableau and Power BI tracking corporate KPIs.
- Wrote complex SQL queries involving window functions, CTEs, and multidimensional joins.
- Built automated ETL data pipelines in Python using Pandas to aggregate financial records.

Data Analyst | Metro Insights (2019 – 2021)
- Conducted exploratory data analysis on consumer behavior datasets using Python and Excel.
- Built and maintained automated weekly reporting suites in Power BI and PostgreSQL.

EDUCATION:
- M.S. in Applied Statistics, Northwestern University (2019)
- B.S. in Economics & Statistics, University of Illinois (2017)
"""
    },
    {
        "candidate_id": "CAND-006",
        "name": "Priya Patel",
        "file_name": "candidate_06_junior_data_scientist.docx",
        "file_type": "DOCX",
        "primary_domain": "Data Science / ML",
        "years_experience": 2,
        "text": """Priya Patel
Junior Data Scientist
Email: priya.patel.ds@example.com | Location: New York, NY

SUMMARY:
Enthusiastic Data Scientist with 2 years of industry experience applying machine learning algorithms, statistical methods, and data visualization techniques. Proficient in Python, Scikit-Learn, Pandas, SQL, and basic NLP concepts (TF-IDF, tokenization).

TECHNICAL SKILLS:
- Machine Learning: Scikit-Learn, Linear Regression, Decision Trees, K-Means, TF-IDF
- Languages & Tools: Python, SQL, Pandas, NumPy, Matplotlib, Seaborn, Jupyter Notebooks
- Development: Git, Basic Docker, Pytest, REST API consumption
- Databases: PostgreSQL, MySQL

EXPERIENCE:
Junior Data Scientist | Quantum Retail (2022 – Present)
- Built predictive customer segmentation models using Scikit-Learn and Python.
- Performed text analysis and TF-IDF feature extraction on customer feedback reviews.
- Extracted and cleaned tabular data using SQL and Pandas.

Data Science Intern | DataSphere Inc (2021 – 2022)
- Assisted in building exploratory data analysis scripts and machine learning models in Python.
- Created data visualization dashboards for marketing performance evaluation.

EDUCATION:
- B.S. in Data Science, New York University (NYU), 2021
"""
    },
    {
        "candidate_id": "CAND-007",
        "name": "Lucas Meyer",
        "file_name": "candidate_07_devops_cloud_engineer.pdf",
        "file_type": "PDF",
        "primary_domain": "DevOps & Cloud Infrastructure",
        "years_experience": 6,
        "text": """Lucas Meyer
Senior DevOps & Cloud Infrastructure Engineer
Email: lucas.meyer.cloud@example.com | Location: Denver, CO

PROFESSIONAL SUMMARY:
DevOps Engineer with 6 years of experience automating cloud infrastructure, orchestrating containers, and maintaining robust CI/CD deployment pipelines. Proficient in AWS, Docker, Kubernetes, Terraform, Python scripting, and Linux system administration.

TECHNICAL EXPERTISE:
- Cloud Platforms: AWS (EC2, S3, ECS, EKS, RDS, IAM, CloudWatch)
- Containerization & Orchestration: Docker, Kubernetes, Helm
- CI/CD & Automation: GitHub Actions, Jenkins, GitLab CI, Terraform, Ansible
- Scripting & Languages: Python, Bash, Shell scripting, Go
- Monitoring & OS: Prometheus, Grafana, Linux (Ubuntu, CentOS), Networking, Git

EXPERIENCE:
Senior DevOps Engineer | CloudScale Inc (2021 – Present)
- Automated multi-region AWS cloud infrastructure using Terraform and Ansible.
- Managed Kubernetes clusters running microservices and automated CI/CD deployments.
- Implemented system monitoring and alerting with Prometheus and Grafana.

Infrastructure Engineer | NetCore Solutions (2018 – 2021)
- Maintained Linux servers and Docker container clusters for production environments.
- Wrote Python automation scripts for data backups, log rotation, and server provisioning.

EDUCATION:
- B.S. in Computer Engineering, University of Colorado Boulder (2018)
"""
    },
    {
        "candidate_id": "CAND-008",
        "name": "Chloe Martin",
        "file_name": "candidate_08_frontend_developer.docx",
        "file_type": "DOCX",
        "primary_domain": "Frontend Web Development",
        "years_experience": 4,
        "text": """Chloe Martin
Frontend Web Developer & UI/UX Specialist
Email: chloe.martin.ui@example.com | Portfolio: chloemartin.design | Location: Los Angeles, CA

SUMMARY:
Creative Frontend Developer with 4 years of experience building modern, responsive, and accessible web interfaces. Highly proficient in HTML5, CSS3, JavaScript, TypeScript, React, and CSS animation frameworks.

CORE SKILLS:
- Web Technologies: HTML5, CSS3, JavaScript (ES6+), TypeScript, SASS, CSS Grid, Flexbox
- Frameworks & Libraries: React, Tailwind CSS, Bootstrap, Webpack, Vite
- Design & UX: Figma, Adobe XD, Responsive Web Design, WCAG Accessibility, Wireframing
- Version Control & Tools: Git, GitHub, NPM, Chrome DevTools, Cross-Browser Testing

EXPERIENCE:
Frontend Developer | PixelCraft Media (2022 – Present)
- Built responsive client websites and web applications using HTML5, modern CSS3, and JavaScript.
- Designed pixel-perfect layouts matching Figma mockups with high cross-device compatibility.
- Improved frontend load times and SEO scores through asset optimization.

UI/UX Web Designer | CreativeStudio (2020 – 2022)
- Created interactive website prototypes, wireframes, and design style guides.
- Developed landing pages using semantic HTML, CSS animations, and vanilla JavaScript.

EDUCATION:
- B.A. in Interactive Media & Web Design, UCLA (2020)
"""
    },
    {
        "candidate_id": "CAND-009",
        "name": "James Wilson",
        "file_name": "candidate_09_digital_marketing_specialist.pdf",
        "file_type": "PDF",
        "primary_domain": "Digital Marketing & Growth",
        "years_experience": 5,
        "text": """James Wilson
Digital Marketing Strategist & SEO Lead
Email: james.wilson.mkt@example.com | Location: Atlanta, GA

PROFESSIONAL SUMMARY:
Results-driven Digital Marketing Strategist with 5+ years of experience leading SEO, content marketing, search engine advertising, and multi-channel customer acquisition campaigns.

CORE COMPETENCIES:
- Search Engine Optimization: Technical SEO, Keyword Research, Link Building, Google Analytics, SEMrush
- Paid Advertising: Google Ads, Meta Ads Manager, LinkedIn Ads, PPC Campaign Management
- Content & Social: Content Strategy, Email Marketing (HubSpot, Mailchimp), Copywriting, Social Media Marketing
- Growth Strategy: Conversion Rate Optimization (CRO), A/B Testing, Lead Generation, Brand Awareness

EXPERIENCE:
Digital Marketing Manager | BrightGrowth Agency (2021 – Present)
- Managed $500K annual advertising budgets across Google Ads and paid social platforms.
- Boosted organic website traffic by 140% through targeted keyword SEO strategies and content optimization.
- Conducted multivariate A/B testing on landing page funnels to optimize conversion rates.

Marketing Specialist | Vanguard Brands (2019 – 2021)
- Created email marketing campaigns reaching 200,000+ subscribers with 28% open rates.
- Coordinated influencer marketing partnerships and social media content calendars.

EDUCATION:
- B.S. in Marketing & Communications, University of Georgia (2019)
"""
    },
    {
        "candidate_id": "CAND-010",
        "name": "Olivia Taylor",
        "file_name": "candidate_10_certified_public_accountant.txt",
        "file_type": "TXT",
        "primary_domain": "Accounting & Finance",
        "years_experience": 7,
        "text": """Olivia Taylor, CPA
Certified Public Accountant & Senior Financial Auditor
Email: olivia.taylor.cpa@example.com | Location: Dallas, TX

SUMMARY:
Licensed Certified Public Accountant (CPA) with 7 years of expertise in corporate financial accounting, regulatory tax compliance, financial reporting, and external audit management.

CORE SKILLS:
- Accounting Standards: US GAAP, IFRS, Internal Controls, SOX Compliance
- Financial Management: General Ledger, Financial Statement Preparation, Balance Sheet Reconciliation, Cash Flow
- Software & Systems: QuickBooks Enterprise, NetSuite ERP, Microsoft Excel (Advanced Modeling), SAP
- Taxation & Audit: Corporate Tax Filing, Payroll Tax, Forensic Auditing, Risk Assessment

EXPERIENCE:
Senior Corporate Accountant | Pinnacle Holdings (2020 – Present)
- Prepared monthly, quarterly, and annual financial statements compliant with US GAAP.
- Managed general ledger reconciliations and automated journal entries using NetSuite ERP.
- Facilitated external audit processes, providing comprehensive financial workpapers.

Staff Auditor | Ernst & Young LLP (2017 – 2020)
- Performed financial statement audits for mid-market manufacturing and technology clients.
- Assessed internal controls and documented audit testing results.

EDUCATION & LICENSING:
- Certified Public Accountant (CPA) License #48291, Texas State Board
- M.S. in Accounting, Southern Methodist University (2017)
- B.B.A. in Accounting, Texas A&M University (2016)
"""
    },
    {
        "candidate_id": "CAND-011",
        "name": "Ethan Davis",
        "file_name": "candidate_11_technical_writer.docx",
        "file_type": "DOCX",
        "primary_domain": "Technical Writing & Documentation",
        "years_experience": 4,
        "text": """Ethan Davis
Senior Technical Writer & Documentation Specialist
Email: ethan.davis.docs@example.com | Location: Portland, OR

SUMMARY:
Technical Writer with 4 years of experience creating clear, developer-friendly software documentation, user manuals, API reference guides, and release notes. Skilled in Markdown, Git, Docusaurus, and cross-functional engineering collaboration.

SKILLS:
- Documentation: Technical Writing, API Documentation, Software Manuals, SDK Guides, Release Notes
- Tools & Markup: Markdown, Git, GitHub, Docusaurus, Postman, Swagger/OpenAPI, Notion, Confluence
- Communication: Information Architecture, Content Strategy, Editing, Proofreading, Developer Relations

EXPERIENCE:
Technical Writer | DevSphere Systems (2022 – Present)
- Authored comprehensive developer documentation and REST API reference guides using Markdown and Swagger.
- Collaborated with engineering teams to document software architecture and installation steps.
- Maintained documentation version control across software releases using Git.

Content Specialist | DocuMedia Corp (2020 – 2022)
- Created user onboarding tutorials, knowledge base FAQs, and how-to guides.

EDUCATION:
- B.A. in English & Technical Communication, University of Oregon (2020)
"""
    },
    {
        "candidate_id": "CAND-012",
        "name": "Sophia Martinez",
        "file_name": "candidate_12_database_administrator.pdf",
        "file_type": "PDF",
        "primary_domain": "Database Administration & SQL",
        "years_experience": 6,
        "text": """Sophia Martinez
Senior Database Administrator & SQL Specialist
Email: sophia.martinez.db@example.com | Location: Phoenix, AZ

PROFESSIONAL SUMMARY:
Database Administrator with 6+ years of experience optimizing relational databases, designing high-availability database architectures, tuning complex SQL queries, and automating database backups and migrations.

AREAS OF EXPERTISE:
- Databases: PostgreSQL, MySQL, Oracle DB, SQL Server, Redis
- Database Operations: Query Optimization, Index Tuning, Replication, High Availability, Disaster Recovery, Backup
- Scripting & Systems: SQL, PL/pgSQL, Python, Bash Scripting, Linux System Administration
- Data Engineering: Schema Design, ETL Data Pipelines, Data Migration, Database Security

EXPERIENCE:
Lead Database Administrator | DataVault Tech (2021 – Present)
- Administered 50+ enterprise PostgreSQL and MySQL production database clusters.
- Optimized query execution plans, indexing strategies, and table partitioning.
- Automated daily backup procedures and failover routines using Python and Bash scripts.

Database Analyst | CoreTech Systems (2018 – 2021)
- Monitored database health, query latency, and connection pooling.
- Designed relational schemas and managed data migrations across staging and production environments.

EDUCATION:
- B.S. in Information Technology, Arizona State University (2018)
"""
    }
]

# -------------------------------------------------------------
# 3. Ground Truth Relevance Labels
# Scale:
# 3 = Highly Relevant (Strong direct match for primary skills and domain)
# 2 = Partially Relevant (Transferable skills, related domain or adjacent stack)
# 1 = Weakly Relevant (Few overlapping general terms, different primary focus)
# 0 = Irrelevant (Completely different profession/domain)
#
# Binary Relevance:
# True (1) if Grade >= 2, False (0) if Grade < 2
# -------------------------------------------------------------
GROUND_TRUTH = [
    # JOB-01: Senior AI/ML Engineer
    {"job_id": "JOB-01", "candidate_id": "CAND-001", "grade": 3, "is_relevant": 1, "rationale": "Direct senior ML/NLP background, PyTorch, FastAPI, AWS, MLOps, TF-IDF."},
    {"job_id": "JOB-01", "candidate_id": "CAND-002", "grade": 3, "is_relevant": 1, "rationale": "Strong NLP data scientist, Scikit-Learn, PyTorch, TF-IDF, text vectorization."},
    {"job_id": "JOB-01", "candidate_id": "CAND-006", "grade": 2, "is_relevant": 1, "rationale": "Junior Data Scientist with ML, Python, Scikit-Learn, TF-IDF background."},
    {"job_id": "JOB-01", "candidate_id": "CAND-003", "grade": 2, "is_relevant": 1, "rationale": "Full-stack Python developer with FastAPI, Docker, AWS, PostgreSQL (solid Python/backend match, lacks deep ML)."},
    {"job_id": "JOB-01", "candidate_id": "CAND-004", "grade": 2, "is_relevant": 1, "rationale": "Backend Python engineer with FastAPI, Docker, CI/CD, SQL (Python/API match, lacks ML/NLP)."},
    {"job_id": "JOB-01", "candidate_id": "CAND-007", "grade": 1, "is_relevant": 0, "rationale": "DevOps engineer with Docker, Kubernetes, AWS, Python (infrastructure match, no ML/NLP)."},
    {"job_id": "JOB-01", "candidate_id": "CAND-005", "grade": 1, "is_relevant": 0, "rationale": "Senior Data Analyst with Python/Pandas/SQL (analytics background, no ML/NLP engineering)."},
    {"job_id": "JOB-01", "candidate_id": "CAND-012", "grade": 1, "is_relevant": 0, "rationale": "Database admin with SQL, PostgreSQL, Python (database match, no ML/NLP)."},
    {"job_id": "JOB-01", "candidate_id": "CAND-008", "grade": 0, "is_relevant": 0, "rationale": "Frontend web developer (React/CSS/HTML), no ML/NLP/Python."},
    {"job_id": "JOB-01", "candidate_id": "CAND-011", "grade": 0, "is_relevant": 0, "rationale": "Technical writer, documentation only."},
    {"job_id": "JOB-01", "candidate_id": "CAND-009", "grade": 0, "is_relevant": 0, "rationale": "Digital marketing & SEO, no software/ML skills."},
    {"job_id": "JOB-01", "candidate_id": "CAND-010", "grade": 0, "is_relevant": 0, "rationale": "Certified Public Accountant, non-technical financial domain."},

    # JOB-02: Full-Stack Python Developer
    {"job_id": "JOB-02", "candidate_id": "CAND-003", "grade": 3, "is_relevant": 1, "rationale": "Direct full-stack Python/FastAPI/Django/HTML5/CSS3/JavaScript/PostgreSQL match."},
    {"job_id": "JOB-02", "candidate_id": "CAND-004", "grade": 3, "is_relevant": 1, "rationale": "Strong backend Python/FastAPI/Django/PostgreSQL engineer with API expertise."},
    {"job_id": "JOB-02", "candidate_id": "CAND-001", "grade": 2, "is_relevant": 1, "rationale": "Strong Python/FastAPI/Docker/PostgreSQL backend skills, primary focus is ML."},
    {"job_id": "JOB-02", "candidate_id": "CAND-008", "grade": 2, "is_relevant": 1, "rationale": "Frontend developer with strong HTML5, CSS3, JavaScript, UI/UX (covers frontend half of role)."},
    {"job_id": "JOB-02", "candidate_id": "CAND-002", "grade": 2, "is_relevant": 1, "rationale": "Data scientist with Python, FastAPI, Docker, PostgreSQL skills."},
    {"job_id": "JOB-02", "candidate_id": "CAND-007", "grade": 1, "is_relevant": 0, "rationale": "DevOps engineer with Docker, CI/CD, Python scripting, AWS."},
    {"job_id": "JOB-02", "candidate_id": "CAND-012", "grade": 1, "is_relevant": 0, "rationale": "Database admin with PostgreSQL, SQL, Python scripting."},
    {"job_id": "JOB-02", "candidate_id": "CAND-006", "grade": 1, "is_relevant": 0, "rationale": "Junior Data Scientist with basic Python/SQL."},
    {"job_id": "JOB-02", "candidate_id": "CAND-005", "grade": 1, "is_relevant": 0, "rationale": "Data analyst with Python/SQL, lacks web development experience."},
    {"job_id": "JOB-02", "candidate_id": "CAND-011", "grade": 0, "is_relevant": 0, "rationale": "Technical writer, documentation only."},
    {"job_id": "JOB-02", "candidate_id": "CAND-009", "grade": 0, "is_relevant": 0, "rationale": "Digital marketing & SEO."},
    {"job_id": "JOB-02", "candidate_id": "CAND-010", "grade": 0, "is_relevant": 0, "rationale": "CPA Accountant."},

    # JOB-03: Data Analyst & BI Specialist
    {"job_id": "JOB-03", "candidate_id": "CAND-005", "grade": 3, "is_relevant": 1, "rationale": "Direct Senior Data Analyst match with SQL, Tableau, Power BI, Python, Pandas."},
    {"job_id": "JOB-03", "candidate_id": "CAND-006", "grade": 2, "is_relevant": 1, "rationale": "Junior Data Scientist with SQL, Pandas, Matplotlib, exploratory data analysis."},
    {"job_id": "JOB-03", "candidate_id": "CAND-012", "grade": 2, "is_relevant": 1, "rationale": "Senior DBA with advanced SQL, query tuning, PostgreSQL, ETL."},
    {"job_id": "JOB-03", "candidate_id": "CAND-002", "grade": 2, "is_relevant": 1, "rationale": "Data scientist with strong Python, Pandas, NumPy, statistical modeling, SQL."},
    {"job_id": "JOB-03", "candidate_id": "CAND-001", "grade": 1, "is_relevant": 0, "rationale": "Senior ML engineer with Python, SQL, statistics (overqualified/focused on deep learning)."},
    {"job_id": "JOB-03", "candidate_id": "CAND-004", "grade": 1, "is_relevant": 0, "rationale": "Backend engineer with SQL, Python, PostgreSQL."},
    {"job_id": "JOB-03", "candidate_id": "CAND-003", "grade": 1, "is_relevant": 0, "rationale": "Full-stack developer with SQL, PostgreSQL, Python."},
    {"job_id": "JOB-03", "candidate_id": "CAND-007", "grade": 0, "is_relevant": 0, "rationale": "DevOps engineer with infrastructure focus."},
    {"job_id": "JOB-03", "candidate_id": "CAND-008", "grade": 0, "is_relevant": 0, "rationale": "Frontend developer (React/CSS)."},
    {"job_id": "JOB-03", "candidate_id": "CAND-009", "grade": 0, "is_relevant": 0, "rationale": "Digital marketing (analytics mentions SEO/Google Ads, not BI)."},
    {"job_id": "JOB-03", "candidate_id": "CAND-010", "grade": 0, "is_relevant": 0, "rationale": "CPA accountant (financial reporting, not BI data analytics)."},
    {"job_id": "JOB-03", "candidate_id": "CAND-011", "grade": 0, "is_relevant": 0, "rationale": "Technical writer."}
]


def create_pdf_resume(filepath, candidate_data):
    """Creates a formatted PDF resume using ReportLab."""
    doc = SimpleDocTemplate(
        filepath,
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'ResumeTitle',
        parent=styles['Heading1'],
        fontSize=15,
        leading=18,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=4
    )
    heading2_style = ParagraphStyle(
        'ResumeHeading2',
        parent=styles['Heading2'],
        fontSize=11,
        leading=15,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=8,
        spaceAfter=3
    )
    body_style = ParagraphStyle(
        'ResumeBody',
        parent=styles['Normal'],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#334155'),
        spaceAfter=3
    )
    
    story = []
    lines = candidate_data["text"].strip().split("\n")
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 4))
            continue
        
        if i == 0:
            story.append(Paragraph(f"<b>{line}</b>", title_style))
        elif line.isupper() and len(line) < 35:
            story.append(Paragraph(f"<b>{line}</b>", heading2_style))
        else:
            formatted_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(formatted_line, body_style))
            
    doc.build(story)


def create_docx_resume(filepath, candidate_data):
    """Creates a formatted DOCX resume using python-docx."""
    doc = Document()
    lines = candidate_data["text"].strip().split("\n")
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        
        if i == 0:
            doc.add_heading(line, level=0)
        elif line.isupper() and len(line) < 35:
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)
            
    doc.save(filepath)


def create_txt_resume(filepath, candidate_data):
    """Creates a plain text resume."""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(candidate_data["text"].strip())


def main():
    print("Generating Job Descriptions...")
    for filename, jd in JOB_DESCRIPTIONS.items():
        filepath = os.path.join(JD_DIR, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(jd["content"].strip())
        print(f"  Created JD: {filename}")

    print("\nGenerating Candidate Resumes in PDF, DOCX, and TXT formats...")
    for cand in CANDIDATES:
        ext = cand["file_type"].lower()
        filepath = os.path.join(RESUMES_DIR, cand["file_name"])
        
        if ext == "pdf":
            create_pdf_resume(filepath, cand)
        elif ext == "docx":
            create_docx_resume(filepath, cand)
        elif ext == "txt":
            create_txt_resume(filepath, cand)
            
        print(f"  Created [{cand['file_type']}] {cand['file_name']} for {cand['name']}")

    print("\nWriting Candidates Metadata CSV...")
    meta_csv_path = os.path.join(METADATA_DIR, "candidates.csv")
    with open(meta_csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["candidate_id", "name", "file_name", "file_type", "primary_domain", "years_experience"]
        )
        writer.writeheader()
        for cand in CANDIDATES:
            writer.writerow({
                "candidate_id": cand["candidate_id"],
                "name": cand["name"],
                "file_name": cand["file_name"],
                "file_type": cand["file_type"],
                "primary_domain": cand["primary_domain"],
                "years_experience": cand["years_experience"]
            })
    print(f"  Saved metadata to {meta_csv_path}")

    print("\nWriting Ground Truth Evaluation Labels CSV...")
    eval_csv_path = os.path.join(EVAL_DIR, "relevance_labels.csv")
    with open(eval_csv_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["job_id", "candidate_id", "grade", "is_relevant", "rationale"]
        )
        writer.writeheader()
        for item in GROUND_TRUTH:
            writer.writerow(item)
    print(f"  Saved ground-truth labels to {eval_csv_path}")

    print("\nSample dataset generation completed successfully!")


if __name__ == "__main__":
    main()
